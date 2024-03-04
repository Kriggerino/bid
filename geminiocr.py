import textwrap
import google.generativeai as genai
from IPython.display import display
from IPython.display import Markdown
from PIL import Image
import os
from dotenv import load_dotenv
from pdf2image import convert_from_path
from docx import Document

load_dotenv()
# Render pdf pages to images
images_from_path = convert_from_path(
    "./output.pdf", output_folder="./result", fmt="jpeg"
)

# Or use `os.getenv('GOOGLE_API_KEY')` to fetch an environment variable.
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)


def delete_all_files_in_directory(directory):
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                delete_all_files_in_directory(file_path)
                os.rmdir(file_path)
        except Exception as e:
            print(f"Failed to delete {file_path}. Reason: {e}")


def to_markdown(text):
    text = text.replace("•", "  *")
    return Markdown(textwrap.indent(text, "> ", predicate=lambda _: True))


# Reading images inside target directory


def process_images_in_directory(directory_path):
    # Initialize GenerativeModel
    model = genai.GenerativeModel("gemini-1.0-pro-vision-latest")
    # Define a list to hold raw text
    raw_text = []

    # Function to process each image
    def process_image(image_path):
        image = Image.open(image_path)
        response = model.generate_content(
            [
                "Return the text exactly as you perceive from the image, any numbered list should be formatted with markdown headings (###), ignore any and all logos or icons present in the image and do not include them in the result returned.",
                image,
            ],
            stream=True,
        )
        response.resolve()
        raw_text.append(response.text)

    # Process each image sequentially
    for filename in os.listdir(directory_path):
        image_path = os.path.join(directory_path, filename)
        process_image(image_path)

    return raw_text


def document_create(raw_text):
    doc = Document()

    for pages in raw_text:
        sections = pages.split("###")
        for section in sections:
            if section.strip():
                # Split section into header and content
                header, content = section.split("\n", 1)
                # Add header to document
                doc.add_heading(header.strip(), level=1)
                # Add content to document as paragraphs
                for paragraph in content.strip().split("\n"):
                    doc.add_paragraph(paragraph.strip())

    doc.save("full_doc.docx")
