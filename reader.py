from time import sleep
from langchain_community.vectorstores import Chroma
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_openai import OpenAIEmbeddings
from langchain.schema import Document as LangchainDocument
from metadata import document_transformer
from langchain_openai import ChatOpenAI
from serpapi import GoogleSearch
import requests
import fitz
from docx.api import Document
from io import BytesIO
import openai
import os
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")
client = openai.OpenAI()


# Insert docx file path for req summary


def get_docx_data(validation_path):
    data = []

    def find_largest_table(doc_path):
        """Find the largest table in a Word document."""
        # Load the Word document
        doc = Document(doc_path)

        # Initialize variables to track the largest table
        largest_table = None
        max_size = 0

        # Iterate through all tables in the document
        for table in doc.tables:
            table_size = len(table.rows) * max(len(row.cells)
                                               for row in table.rows)
            if table_size > max_size:
                max_size = table_size
                largest_table = table

        return largest_table

    table = find_largest_table(validation_path)
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data


def get_requirements(docx_data):
    try:
        response = client.chat.completions.create(
            model="gpt-4-1106-preview",
            messages=[{
                "role": "system",
                "content": "You will be provided with a representation of a technical requirement table written in Vietnamese as a Python list containing many dictionaries in the following format: data=[{...}, {...}, ...]. Return a shortened requirement list in Vietnamese listing the minimum models, the amounts needed, the specs related to them , additional maintenance requirements (if any) from the data provided. If a component name is mentioned in the same sentence as the phrase 'Chúng tôi chào...', ignore all instances of the component name and do not return it in the list. Each item in the list should be separated by a new line with '------'. Do not introduce yourself or explain. Only use the data provided to you. Do not use unnecessary whitespace. Do not return the raw data in response."
            }, {
                "role": "user",
                "content": f"Here is the provided data: data=```{docx_data}``` "
            }],
            temperature=0.1,
            top_p=1,

        )
    except Exception as e:
        print(e)

    return response.choices[0].message.content


def get_datasheet_name(docx_data):
    try:
        respond2 = client.chat.completions.create(
            model="gpt-3.5-turbo-16k",
            messages=[{
                "role": "system",
                "content": "You will be provided with a representation of a technical requirement table written in Vietnamese as a Python list containing many dictionaries in the following format: data=[{...}, {...}, ...]. Return a list containing the name of datasheet models pdfs that is referred to inside double air quotes. Each item in the final list should be seperated by a new line. If there's multiple instances of the same element in the list, only return 1 instance in the list. Do not introduce yourself or explain. Only use the data provided to you. Do not use unnecessary whitespace."
            }, {
                "role": "user",
                "content": f"Here is the provided data: data=```{docx_data}``` "
            }],
            temperature=0.2,
            top_p=1,
        )
    except Exception as e:
        print(e)

    arr = respond2.choices[0].message.content.split("\n")
    cleaned_arr = [name.replace('.pdf', '').strip('"') for name in arr]
    return cleaned_arr


def get_datasheets(arr):
    downloaded_files = []
    links = []
    for name in arr:
        search = GoogleSearch({
            "engine": "google",
            "q": f"{name} filetype:pdf",
            "api_key": f'{os.getenv("SERPAPI_KEY")}',
            "num": 1,
        })
        result = search.get_dict()
        links.append(result.get('organic_results')[0]['link'])

    for url in links:
        try:
            response = requests.get(url, allow_redirects=True)
            # Extracting the filename from the URL
            max_size = 32* 1024 * 1024
            filename = url.split("/")[-1]
            file_size = len(response.content)
            if file_size < max_size:
                with open(f"./documents/pdfs/{filename}", 'wb') as file:
                    file.write(response.content)
                # Append the filename to the global list
                downloaded_files.append(f"./documents/pdfs/{filename}")
                print(f"Downloaded: {filename}")
            else:
                print(f"Skipped {filename}: File size exceeds maximum size limit.")
        except Exception as e:
            print(f"Error downloading {url}: {e}")

    return downloaded_files


def read_datasheets_specs(downloaded_files):
    chatpdf_sources = []
    results = []
    for file_paths in downloaded_files:
        try:
            files = [('file', ('file', open(f'{file_paths}','rb'), 'application/octet-stream'))]
            headers = {
                'x-api-key': 'sec_cZErkSMw8uyphfqgmOoY7rL0QHzX3ky4'
            }
            response = requests.post(
                'https://api.chatpdf.com/v1/sources/add-file', headers=headers, files=files)

            if response.status_code == 200:
                chatpdf_sources.append(response.json()['sourceId'])
            else:
                print('Status:', response.status_code)
                print('Error:', response.text)
        except Exception as e:
            print(e)
    for source in chatpdf_sources:
        try:
            headers = {
                'x-api-key': 'sec_cZErkSMw8uyphfqgmOoY7rL0QHzX3ky4',
                "Content-Type": "application/json",
            }

            data = {
                'sourceId': f'{source}',
                'messages': [
                    {
                        'role': "user",
                        'content': "List all the features and/or specifications mentioned in the document. If there are multiple devices mentioned, list both devices seperately. Go into detail for each specifications. Always include the key specs such as CPU, RAM, Storage, Power,... if it exists in the document. Do not introduce yourself or explain. Always open the result with 'The features and specification of the <name of device(s)> is:... '",
                    }
                ]
            }
            response = requests.post(
                'https://api.chatpdf.com/v1/chats/message', headers=headers, json=data)
            print(response)

            if response.status_code == 200:
                results.append(response.json()['content'])
            else:
                print('Status:', response.status_code)
                print('Error:', response.text)
        except Exception as e:
            print(e)
            return []

    return results


def compare(datasheets_specs, requirements, query):
    documents = [LangchainDocument(page_content=result) for result in datasheets_specs]
    # Creating context for vector stores
    enhanced_documents = document_transformer.transform_documents(documents)
    enhanced_documents.append(
        LangchainDocument(page_content=requirements, metadata={"device": "requirements"}))
    # Creating vector store
    embeddings = OpenAIEmbeddings()
    db = Chroma.from_documents(enhanced_documents, embeddings)
    # Retrieving results
    retriever = db.as_retriever(search_type="mmr")
    # Creating chain for generation
    llm = ChatOpenAI(model_name="gpt-4-1106-preview", temperature=0)

    template = """Use the following pieces of context to answer the question at the end. Do not introduce or explain yourself. Keep the answer short and concise.

    {context}

    Question: {question}

    Answer:"""
    custom_rag_prompt = PromptTemplate.from_template(template)

    chain = create_stuff_documents_chain(llm, custom_rag_prompt)

    return chain.invoke({"context": enhanced_documents, "question": query})



# Workflow tests, replicate on API call
# requirements = get_requirements(data)

# datasheets = get_datasheet_name(data)


# downloaded_files = get_datasheets(datasheets)



# datasheets_specs = read_datasheets_specs(downloaded_files)
# print(datasheets_specs)
# final = compare(datasheets_specs=datasheets_specs, requirements=requirements,
#                 query="Compare the 'Quantum 6600 Security Gateway' to the requirements of 'Thiết bị tường lửa cho phân vùng kết nối Cloud' and answer if they match. If any specs doesn't match, list how they do not match.")

# print(final)
